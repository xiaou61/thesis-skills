#!/usr/bin/env python3
"""Conservatively repair a thesis .docx using an extracted template profile."""

from __future__ import annotations

import argparse
import json
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt
import yaml


@dataclass
class RepairAction:
    location: str
    change: str
    before: str
    after: str
    status: str


def normalize_text(value: str) -> str:
    return " ".join(value.split()).strip()


def extract_story_text(paragraphs: Any) -> str:
    return normalize_text(" ".join(p.text for p in paragraphs if getattr(p, "text", "").strip()))


def set_story_text(paragraphs: Any, text: str) -> tuple[str, str]:
    before = extract_story_text(paragraphs)
    if not paragraphs:
        return before, before
    paragraphs[0].text = text
    for paragraph in paragraphs[1:]:
        paragraph.text = ""
    after = extract_story_text(paragraphs)
    return before, after


def cm_repr(value: Any) -> str:
    try:
        return f"{round(value.cm, 2)}cm"
    except Exception:
        return str(value)


def apply_margin(section: Any, attr: str, expected_cm: float | None, actions: list[RepairAction], location: str) -> None:
    if expected_cm is None:
        return
    current = getattr(section, attr)
    before = cm_repr(current)
    setattr(section, attr, Cm(expected_cm))
    after = cm_repr(getattr(section, attr))
    if before != after:
        actions.append(
            RepairAction(location, f"set {attr}", before, after, "done")
        )


def apply_orientation(section: Any, orientation: str | None, actions: list[RepairAction], location: str) -> None:
    if not orientation:
        return
    current = "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait"
    target = orientation.lower()
    if current == target:
        return
    before = current
    if target == "landscape":
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
    elif target == "portrait":
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = new_width
        section.page_height = new_height
    else:
        return
    after = "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait"
    actions.append(RepairAction(location, "set orientation", before, after, "done"))


def first_effective_text(items: list[dict[str, Any]]) -> str:
    texts = []
    for item in items:
        if not isinstance(item, dict):
            continue
        texts.extend(item.get("text", []))
    return normalize_text(" ".join(texts))


def alignment_from_string(value: str | None) -> Any:
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get((value or "").lower())


def alignment_repr(value: Any) -> str:
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        None: "None",
    }
    return mapping.get(value, str(value))


def set_style_font_name(style: Any, font_name: str) -> None:
    style.font.name = font_name
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = rpr._add_rFonts()
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def apply_line_spacing_value(paragraph_format: Any, expected_pt: float, expected_rule: str | None) -> None:
    rule = (expected_rule or "").lower()
    if rule == "auto":
        paragraph_format.line_spacing = round(float(expected_pt) / 12.0, 3)
    else:
        paragraph_format.line_spacing = Pt(expected_pt)


def set_spacing_xml(owner_element: Any, expected_pt: float, expected_rule: str | None) -> None:
    ppr = owner_element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        owner_element.append(ppr)
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:line"), str(int(round(float(expected_pt) * 20))))
    if expected_rule:
        spacing.set(qn("w:lineRule"), expected_rule)


def apply_style_property(actions: list[RepairAction], location: str, label: str, before: str, after: str) -> None:
    if before != after:
        actions.append(RepairAction(location, label, before, after, "done"))


def collect_style_targets(template_profile: dict[str, Any], overrides: dict[str, Any] | None) -> list[str]:
    override_root = overrides.get("template_rule_overrides", {}) if isinstance(overrides, dict) else {}
    targets: list[str] = []
    for key in ("heading_style_map", "body_style_candidates", "reference_style_candidates", "caption_style_candidates"):
        for item in override_root.get(key, []) or []:
            if isinstance(item, dict):
                style_id = str(item.get("style_id", "")).strip()
                if style_id and style_id not in targets:
                    targets.append(style_id)
    if "Normal" not in targets and any(item.get("style_id") == "Normal" for item in template_profile.get("styles", [])):
        targets.append("Normal")
    return targets


def find_style_by_id(document: Document, style_id: str) -> Any | None:
    for style in document.styles:
        if getattr(style, "style_id", None) == style_id:
            return style
    return None


def story_for_type(section: Any, kind: str, ref_type: str) -> Any | None:
    ref_type = (ref_type or "default").lower()
    if kind == "header":
        if ref_type == "first":
            return section.first_page_header
        if ref_type == "even":
            return section.even_page_header
        return section.header
    if kind == "footer":
        if ref_type == "first":
            return section.first_page_footer
        if ref_type == "even":
            return section.even_page_footer
        return section.footer
    return None


def repair_style_formats(document: Document, template_profile: dict[str, Any], overrides: dict[str, Any] | None, actions: list[RepairAction]) -> None:
    style_index = {
        item.get("style_id"): item
        for item in template_profile.get("styles", [])
        if isinstance(item, dict) and item.get("style_id")
    }
    for style_id in collect_style_targets(template_profile, overrides):
        template_style = style_index.get(style_id)
        if not template_style:
            continue
        style = find_style_by_id(document, style_id)
        if style is None:
            actions.append(RepairAction(f"style {style_id}", "style definition", "missing", "missing", "skipped_manual_review_required"))
            continue

        location = f"style {template_style.get('name') or style_id} ({style_id})"
        expected_font = template_style.get("font_east_asia") or template_style.get("font_ascii")
        if expected_font:
            before = style.font.name or ""
            set_style_font_name(style, expected_font)
            after = style.font.name or ""
            apply_style_property(actions, location, "set font family", before or "(empty)", after or "(empty)")

        expected_size = template_style.get("font_size_pt")
        if expected_size is not None:
            before = str(round(style.font.size.pt, 1)) if style.font.size else "(empty)"
            style.font.size = Pt(expected_size)
            after = str(round(style.font.size.pt, 1)) if style.font.size else "(empty)"
            apply_style_property(actions, location, "set font size", before, after)

        expected_bold = template_style.get("bold")
        if expected_bold is not None:
            before = str(style.font.bold)
            style.font.bold = bool(expected_bold)
            after = str(style.font.bold)
            apply_style_property(actions, location, "set bold", before, after)

        expected_italic = template_style.get("italic")
        if expected_italic is not None:
            before = str(style.font.italic)
            style.font.italic = bool(expected_italic)
            after = str(style.font.italic)
            apply_style_property(actions, location, "set italic", before, after)

        pf = style.paragraph_format
        expected_align = alignment_from_string(template_style.get("alignment"))
        if expected_align is not None:
            before = alignment_repr(pf.alignment)
            pf.alignment = expected_align
            after = alignment_repr(pf.alignment)
            apply_style_property(actions, location, "set alignment", before, after)

        if template_style.get("spacing_before_pt") is not None:
            before = str(round(pf.space_before.pt, 1)) if pf.space_before else "(empty)"
            pf.space_before = Pt(template_style["spacing_before_pt"])
            after = str(round(pf.space_before.pt, 1)) if pf.space_before else "(empty)"
            apply_style_property(actions, location, "set spacing before", before, after)

        if template_style.get("spacing_after_pt") is not None:
            before = str(round(pf.space_after.pt, 1)) if pf.space_after else "(empty)"
            pf.space_after = Pt(template_style["spacing_after_pt"])
            after = str(round(pf.space_after.pt, 1)) if pf.space_after else "(empty)"
            apply_style_property(actions, location, "set spacing after", before, after)

        if template_style.get("line_spacing_pt") is not None:
            before = str(round(pf.line_spacing.pt, 1)) if hasattr(pf.line_spacing, "pt") else str(pf.line_spacing)
            apply_line_spacing_value(pf, float(template_style["line_spacing_pt"]), template_style.get("line_spacing_rule"))
            set_spacing_xml(style._element, float(template_style["line_spacing_pt"]), template_style.get("line_spacing_rule"))
            after = str(round(pf.line_spacing.pt, 1)) if hasattr(pf.line_spacing, "pt") else str(pf.line_spacing)
            apply_style_property(actions, location, "set line spacing", before, after)

        if template_style.get("first_line_indent_cm") is not None:
            before = f"{round(pf.first_line_indent.cm, 2)}cm" if pf.first_line_indent else "(empty)"
            pf.first_line_indent = Cm(template_style["first_line_indent_cm"])
            after = f"{round(pf.first_line_indent.cm, 2)}cm" if pf.first_line_indent else "(empty)"
            apply_style_property(actions, location, "set first-line indent", before, after)

        if template_style.get("left_indent_cm") is not None:
            before = f"{round(pf.left_indent.cm, 2)}cm" if pf.left_indent else "(empty)"
            pf.left_indent = Cm(template_style["left_indent_cm"])
            after = f"{round(pf.left_indent.cm, 2)}cm" if pf.left_indent else "(empty)"
            apply_style_property(actions, location, "set left indent", before, after)

        if template_style.get("hanging_indent_cm") is not None:
            before = f"{round(abs(pf.first_line_indent.cm), 2)}cm" if pf.first_line_indent and pf.first_line_indent.cm < 0 else "(empty)"
            pf.first_line_indent = Cm(-abs(float(template_style["hanging_indent_cm"])))
            after = f"{round(abs(pf.first_line_indent.cm), 2)}cm" if pf.first_line_indent and pf.first_line_indent.cm < 0 else "(empty)"
            apply_style_property(actions, location, "set hanging indent", before, after)

        for paragraph in document.paragraphs:
            style_obj = getattr(paragraph, "style", None)
            if getattr(style_obj, "style_id", None) != style_id:
                continue
            plocation = f"{location} paragraph '{paragraph.text[:20]}'"
            pformat = paragraph.paragraph_format
            expected_align = alignment_from_string(template_style.get("alignment"))
            if expected_align is not None and pformat.alignment is not None and pformat.alignment != expected_align:
                before = alignment_repr(pformat.alignment)
                pformat.alignment = expected_align
                after = alignment_repr(pformat.alignment)
                apply_style_property(actions, plocation, "set paragraph alignment override", before, after)
            if template_style.get("first_line_indent_cm") is not None and pformat.first_line_indent is not None:
                before_val = round(pformat.first_line_indent.cm, 2)
                expected_val = round(float(template_style["first_line_indent_cm"]), 2)
                if before_val != expected_val:
                    before = f"{before_val}cm"
                    pformat.first_line_indent = Cm(template_style["first_line_indent_cm"])
                    after = f"{round(pformat.first_line_indent.cm, 2)}cm" if pformat.first_line_indent else "(empty)"
                    apply_style_property(actions, plocation, "set paragraph first-line indent override", before, after)
            if template_style.get("left_indent_cm") is not None and pformat.left_indent is not None:
                before_val = round(pformat.left_indent.cm, 2)
                expected_val = round(float(template_style["left_indent_cm"]), 2)
                if before_val != expected_val:
                    before = f"{before_val}cm"
                    pformat.left_indent = Cm(template_style["left_indent_cm"])
                    after = f"{round(pformat.left_indent.cm, 2)}cm" if pformat.left_indent else "(empty)"
                    apply_style_property(actions, plocation, "set paragraph left-indent override", before, after)
            if template_style.get("hanging_indent_cm") is not None and pformat.first_line_indent is not None and pformat.first_line_indent.cm < 0:
                before_val = round(abs(pformat.first_line_indent.cm), 2)
                expected_val = round(float(template_style["hanging_indent_cm"]), 2)
                if before_val != expected_val:
                    before = f"{before_val}cm"
                    pformat.first_line_indent = Cm(-abs(float(template_style["hanging_indent_cm"])))
                    after = f"{round(abs(pformat.first_line_indent.cm), 2)}cm" if pformat.first_line_indent and pformat.first_line_indent.cm < 0 else "(empty)"
                    apply_style_property(actions, plocation, "set paragraph hanging-indent override", before, after)
            if template_style.get("spacing_after_pt") is not None and pformat.space_after is not None:
                before_val = round(pformat.space_after.pt, 1)
                expected_val = round(float(template_style["spacing_after_pt"]), 1)
                if before_val != expected_val:
                    before = f"{before_val}pt"
                    pformat.space_after = Pt(template_style["spacing_after_pt"])
                    after = f"{round(pformat.space_after.pt, 1)}pt" if pformat.space_after else "(empty)"
                    apply_style_property(actions, plocation, "set paragraph spacing-after override", before, after)
            if template_style.get("line_spacing_pt") is not None and hasattr(pformat.line_spacing, "pt"):
                before_val = round(pformat.line_spacing.pt, 1)
                expected_val = round(float(template_style["line_spacing_pt"]), 1)
                if before_val != expected_val:
                    before = f"{before_val}pt"
                    apply_line_spacing_value(pformat, float(template_style["line_spacing_pt"]), template_style.get("line_spacing_rule"))
                    set_spacing_xml(paragraph._element, float(template_style["line_spacing_pt"]), template_style.get("line_spacing_rule"))
                    after = f"{round(pformat.line_spacing.pt, 1)}pt" if hasattr(pformat.line_spacing, "pt") else str(pformat.line_spacing)
                    apply_style_property(actions, plocation, "set paragraph line-spacing override", before, after)


def repair_docx(thesis_docx: Path, template_profile: dict[str, Any], output_docx: Path, overrides: dict[str, Any] | None = None) -> list[RepairAction]:
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(thesis_docx, output_docx)
    document = Document(output_docx)
    actions: list[RepairAction] = []

    template_sections = template_profile.get("sections", [])
    thesis_sections = list(document.sections)

    if len(template_sections) != len(thesis_sections):
        actions.append(
            RepairAction(
                "document",
                "section-count-check",
                str(len(thesis_sections)),
                str(len(template_sections)),
                "skipped_manual_review_required",
            )
        )

    for idx, expected in enumerate(template_sections, 1):
        if idx > len(thesis_sections):
            break
        section = thesis_sections[idx - 1]
        location = f"section {idx}"
        page = expected.get("page", {})
        margins = expected.get("margins_cm", {})
        expected_first_page = bool(expected.get("different_first_page"))

        apply_orientation(section, page.get("orientation"), actions, location)
        before_first_page = str(bool(section.different_first_page_header_footer))
        section.different_first_page_header_footer = expected_first_page
        after_first_page = str(bool(section.different_first_page_header_footer))
        apply_style_property(actions, location, "set different-first-page flag", before_first_page, after_first_page)
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            expected_key = attr.replace("_margin", "")
            apply_margin(section, attr, margins.get(expected_key), actions, location)

        for item in expected.get("headers", []):
            if not isinstance(item, dict):
                continue
            expected_header = first_effective_text([item])
            if expected_header:
                story = story_for_type(section, "header", str(item.get("type", "default")))
                if story is None:
                    continue
                before, after = set_story_text(story.paragraphs, expected_header)
                if before != after:
                    actions.append(
                        RepairAction(
                            location,
                            f"set {item.get('type', 'default')} header text",
                            before or "(empty)",
                            after or "(empty)",
                            "done",
                        )
                    )

        for item in expected.get("footers", []):
            if not isinstance(item, dict):
                continue
            expected_footer = first_effective_text([item])
            if expected_footer:
                story = story_for_type(section, "footer", str(item.get("type", "default")))
                if story is None:
                    continue
                before, after = set_story_text(story.paragraphs, expected_footer)
                if before != after:
                    actions.append(
                        RepairAction(
                            location,
                            f"set {item.get('type', 'default')} footer text",
                            before or "(empty)",
                            after or "(empty)",
                            "done",
                        )
                    )

    repair_style_formats(document, template_profile, overrides, actions)
    document.save(output_docx)
    return actions


def render_markdown(source_docx: Path, output_docx: Path, actions: list[RepairAction]) -> str:
    lines = [
        "# DOCX Template Repair Report",
        "",
        f"- Source thesis: `{source_docx}`",
        f"- Repaired thesis: `{output_docx}`",
        f"- Action count: `{len(actions)}`",
        "",
    ]
    if not actions:
        lines.append("- No repair actions were needed.")
        return "\n".join(lines) + "\n"

    for action in actions:
        lines.append(
            f"- `{action.location}` {action.change}: `{action.before}` -> `{action.after}` ({action.status})"
        )
    lines.append("")
    lines.append("Manual review remains required for TOC fields, page-number fields, cross-references, and complex pagination.")
    return "\n".join(lines) + "\n"


def main() -> int:
    parser = argparse.ArgumentParser(description="Repair a thesis .docx using a template profile.")
    parser.add_argument("thesis_docx", help="Thesis .docx file to repair.")
    parser.add_argument("template_profile_json", help="template-profile.json path extracted from school template.")
    parser.add_argument("--template-rule-overrides", help="Optional template-rule-overrides.yaml path.")
    parser.add_argument("--out-docx", help="Output repaired .docx path.")
    parser.add_argument("--out-report", default="paper-context/template-compare/template-repair-report.md", help="Output repair markdown report.")
    parser.add_argument("--json-out", help="Optional JSON action log path.")
    args = parser.parse_args()

    thesis_docx = Path(args.thesis_docx).resolve()
    template_json = Path(args.template_profile_json).resolve()
    template_profile = json.loads(template_json.read_text(encoding="utf-8"))
    overrides = None
    if args.template_rule_overrides:
        overrides = yaml.safe_load(Path(args.template_rule_overrides).resolve().read_text(encoding="utf-8"))

    out_docx = Path(args.out_docx).resolve() if args.out_docx else thesis_docx.with_name(f"{thesis_docx.stem}_repaired{thesis_docx.suffix}")
    actions = repair_docx(thesis_docx, template_profile, out_docx, overrides)

    report_path = Path(args.out_report).resolve()
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(render_markdown(thesis_docx, out_docx, actions), encoding="utf-8")
    print(f"Wrote {out_docx}")
    print(f"Wrote {report_path}")

    if args.json_out:
        json_path = Path(args.json_out).resolve()
        json_path.parent.mkdir(parents=True, exist_ok=True)
        payload = {
            "source_docx": str(thesis_docx),
            "repaired_docx": str(out_docx),
            "actions": [action.__dict__ for action in actions],
        }
        json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"Wrote {json_path}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
